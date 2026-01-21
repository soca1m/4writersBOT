"""
Утилиты для парсинга файлов (PDF, DOCX, TXT)
"""
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Извлекает текст из PDF файла

    Args:
        file_path: Путь к PDF файлу

    Returns:
        Извлеченный текст
    """
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        text = ""

        for page in reader.pages:
            text += page.extract_text() + "\n"

        logger.info(f"Extracted {len(text)} characters from PDF: {file_path}")
        return text

    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """
    Извлекает текст из DOCX файла

    Args:
        file_path: Путь к DOCX файлу

    Returns:
        Извлеченный текст
    """
    try:
        from docx import Document

        doc = Document(file_path)

        # Извлекаем текст из параграфов
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Извлекаем текст из таблиц
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables_text.append(row_text)

        # Объединяем всё
        all_text = []
        if paragraphs:
            all_text.append("\n".join(paragraphs))
        if tables_text:
            all_text.append("\n=== Tables ===\n" + "\n".join(tables_text))

        text = "\n\n".join(all_text)

        logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
        return text

    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """
    Читает текстовый файл

    Args:
        file_path: Путь к TXT файлу

    Returns:
        Содержимое файла
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        logger.info(f"Read {len(text)} characters from TXT: {file_path}")
        return text

    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {e}")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """
    Автоматически определяет тип файла и извлекает текст

    Args:
        file_path: Путь к файлу

    Returns:
        Извлеченный текст
    """
    path = Path(file_path)

    if not path.exists():
        logger.error(f"File not found: {file_path}")
        return ""

    extension = path.suffix.lower()

    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif extension in ['.txt', '.text']:
        return extract_text_from_txt(file_path)
    else:
        logger.warning(f"Unsupported file type: {extension}")
        return ""


def parse_multiple_files(file_paths: list) -> str:
    """
    Парсит несколько файлов и объединяет их содержимое

    Args:
        file_paths: Список путей к файлам

    Returns:
        Объединенный текст из всех файлов
    """
    all_text = []

    for file_path in file_paths:
        text = extract_text_from_file(file_path)
        if text:
            all_text.append(f"=== File: {Path(file_path).name} ===\n{text}\n")

    return "\n\n".join(all_text)
